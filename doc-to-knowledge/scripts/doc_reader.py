"""飞书文档解析模块 — 读取 docx 并按章节切分。

读取飞书 docx 文档（如视频转录稿），将其按 Heading2 分隔符拆分为多个章节，
供 knowledge_extractor 逐章提取知识点。

支持三种读取方式：
1. 飞书在线文档 → document blocks API
2. 飞书云盘上传的 .docx 文件 → 下载后用 python-docx 本地解析
3. 本地 Obsidian Markdown 文件 → 直接读取解析

文档结构约定（由上游保证）：
├── 标题行（page block）
├── 摘要信息段落（文本 block）
├── 分割线
├── Heading2：章节 1 标题
│   ├── 文本段落 ...
│   └── 分割线
├── Heading2：章节 2 标题
│   ├── 文本段落 ...
│   └── 分割线
└── ...
"""

from __future__ import annotations

import logging
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from feishu_client import FeishuClient

logger = logging.getLogger(__name__)

# 飞书 docx block_type 常量
BLOCK_TYPE_PAGE = 1       # 文档根节点（page）
BLOCK_TYPE_TEXT = 2       # 文本段落
BLOCK_TYPE_HEADING1 = 3   # heading1
BLOCK_TYPE_HEADING2 = 4   # heading2
BLOCK_TYPE_HEADING3 = 5   # heading3
BLOCK_TYPE_HEADING4 = 6   # heading4
BLOCK_TYPE_HEADING5 = 7   # heading5
BLOCK_TYPE_HEADING6 = 8   # heading6
BLOCK_TYPE_HEADING7 = 9   # heading7
BLOCK_TYPE_HEADING8 = 10  # heading8
BLOCK_TYPE_HEADING9 = 11  # heading9
BLOCK_TYPE_BULLET = 12    # 无序列表
BLOCK_TYPE_ORDERED = 13   # 有序列表
BLOCK_TYPE_CODE = 14      # 代码块
BLOCK_TYPE_QUOTE = 15     # 引用
BLOCK_TYPE_DIVIDER = 22   # 分割线


@dataclass
class Chapter:
    """一个文档章节。"""
    title: str
    paragraphs: list[str] = field(default_factory=list)

    @property
    def content(self) -> str:
        """将所有段落拼接为完整文本。"""
        return "\n\n".join(p for p in self.paragraphs if p.strip())

    @property
    def word_count(self) -> int:
        return len(self.content)


@dataclass
class DocumentContent:
    """解析后的文档内容。"""
    doc_title: str
    doc_token: str
    summary: str
    chapters: list[Chapter]

    @property
    def total_word_count(self) -> int:
        return sum(ch.word_count for ch in self.chapters)

    def to_dict(self) -> dict[str, Any]:
        return {
            "doc_title": self.doc_title,
            "doc_token": self.doc_token,
            "summary": self.summary,
            "total_chapters": len(self.chapters),
            "total_word_count": self.total_word_count,
            "chapters": [
                {
                    "title": ch.title,
                    "word_count": ch.word_count,
                    "paragraph_count": len(ch.paragraphs),
                }
                for ch in self.chapters
            ],
        }


def extract_text_from_block(block: dict[str, Any]) -> str:
    """从一个 block 中提取纯文本内容。

    飞书 block 的文本内容存储在 elements 数组中，
    每个 element 可能是 text_run / mention_user 等类型。
    """
    block_type = block.get("block_type", 0)

    # 根据 block_type 确定内容字段名
    content_key_map = {
        BLOCK_TYPE_TEXT: "text",
        BLOCK_TYPE_HEADING1: "heading1",
        BLOCK_TYPE_HEADING2: "heading2",
        BLOCK_TYPE_HEADING3: "heading3",
        BLOCK_TYPE_HEADING4: "heading4",
        BLOCK_TYPE_HEADING5: "heading5",
        BLOCK_TYPE_HEADING6: "heading6",
        BLOCK_TYPE_HEADING7: "heading7",
        BLOCK_TYPE_HEADING8: "heading8",
        BLOCK_TYPE_HEADING9: "heading9",
        BLOCK_TYPE_BULLET: "bullet",
        BLOCK_TYPE_ORDERED: "ordered",
        BLOCK_TYPE_CODE: "code",
        BLOCK_TYPE_QUOTE: "quote",
    }

    content_key = content_key_map.get(block_type)
    if not content_key:
        return ""

    content_data = block.get(content_key, {})
    elements = content_data.get("elements", [])

    text_parts = []
    for elem in elements:
        text_run = elem.get("text_run", {})
        text = text_run.get("content", "")
        if text:
            text_parts.append(text)

    return "".join(text_parts).strip()


def read_document(client: FeishuClient, doc_token: str) -> DocumentContent:
    """读取飞书文档并按章节切分。

    自动检测文档类型，按优先级 fallback：
    A. 飞书在线文档 → document blocks API
    B. 飞书云盘上传文件 → 下载后 python-docx 本地解析

    Args:
        client: 飞书 API 客户端
        doc_token: 文档 ID / token

    Returns:
        解析后的 DocumentContent，包含标题、摘要和章节列表。
    """
    logger.info("开始读取飞书文档: %s", doc_token)

    # 尝试方法 A：飞书在线文档 API
    try:
        meta = client.get_document_meta(doc_token)
        doc_title = meta.get("title", "未命名文档")
        logger.info("文档标题: %s（飞书在线文档模式）", doc_title)
        return _read_from_blocks(client, doc_token, doc_title)
    except Exception as e:
        logger.info("在线文档 API 不可用 (%s)，尝试下载本地解析...", e)

    # 方法 B：下载 .docx 文件后本地解析
    try:
        return _read_from_downloaded_docx(client, doc_token)
    except Exception as e:
        raise RuntimeError(
            f"所有读取方式均失败 (token={doc_token})。"
            f"建议使用 --local-docx 指定本地 .docx 文件路径。"
            f"错误: {e}"
        ) from e


def _read_from_blocks(client: FeishuClient, doc_token: str, doc_title: str) -> DocumentContent:
    """从飞书在线文档 blocks API 读取并解析。"""
    blocks = client.get_document_blocks(doc_token)
    logger.info("共获取 %d 个 block", len(blocks))

    summary_parts: list[str] = []
    chapters: list[Chapter] = []
    current_chapter: Chapter | None = None
    found_first_heading2 = False

    for block in blocks:
        block_type = block.get("block_type", 0)

        if block_type == BLOCK_TYPE_PAGE:
            continue
        if block_type == BLOCK_TYPE_DIVIDER:
            continue

        if block_type == BLOCK_TYPE_HEADING2:
            found_first_heading2 = True
            heading_text = extract_text_from_block(block)
            clean_title = _clean_chapter_title(heading_text)

            if current_chapter and current_chapter.paragraphs:
                chapters.append(current_chapter)

            current_chapter = Chapter(title=clean_title)
            logger.debug("发现章节: %s", clean_title)
            continue

        text = extract_text_from_block(block)
        if not text:
            continue

        if not found_first_heading2:
            summary_parts.append(text)
        elif current_chapter is not None:
            current_chapter.paragraphs.append(text)

    if current_chapter and current_chapter.paragraphs:
        chapters.append(current_chapter)

    summary = "\n".join(summary_parts).strip()

    logger.info(
        "文档解析完成: %d 个章节, 摘要 %d 字, 正文 %d 字",
        len(chapters), len(summary), sum(ch.word_count for ch in chapters),
    )

    if not chapters:
        logger.warning("未找到 Heading2 章节分隔，将整个文档作为单一章节处理")
        all_text_parts = summary_parts[:]
        for block in blocks:
            block_type = block.get("block_type", 0)
            if block_type in (BLOCK_TYPE_PAGE, BLOCK_TYPE_DIVIDER):
                continue
            text = extract_text_from_block(block)
            if text and text not in summary_parts:
                all_text_parts.append(text)
        if all_text_parts:
            chapters = [Chapter(title=doc_title, paragraphs=all_text_parts)]

    return DocumentContent(
        doc_title=doc_title,
        doc_token=doc_token,
        summary=summary,
        chapters=chapters,
    )


def _read_from_downloaded_docx(client: FeishuClient, file_token: str) -> DocumentContent:
    """下载飞书云盘的 .docx 文件后用 python-docx 本地解析。

    适用于：上传到飞书云盘的 .docx 文件（非飞书原生在线文档）。
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("需要安装 python-docx 库: pip install python-docx")

    # 获取文件名信息
    doc_title = f"文档_{file_token[:8]}"
    try:
        meta = client.get_file_meta(file_token)
        doc_title = meta.get("title", doc_title)
        # 去掉 .docx 扩展名
        if doc_title.endswith(".docx"):
            doc_title = doc_title[:-5]
        logger.info("文件标题: %s（下载本地解析模式）", doc_title)
    except Exception as e:
        logger.warning("获取文件元信息失败: %s，使用默认标题", e)

    # 下载到临时文件
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        client.download_file(file_token, tmp_path)
        docx_doc = DocxDocument(tmp_path)
    except Exception as e:
        # 清理临时文件
        Path(tmp_path).unlink(missing_ok=True)
        raise RuntimeError(f"下载或解析 .docx 文件失败: {e}") from e

    # 解析段落
    summary_parts: list[str] = []
    chapters: list[Chapter] = []
    current_chapter: Chapter | None = None
    found_first_heading2 = False

    for para in docx_doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        # 跳过分割线（常见的分割线样式）
        if text.startswith("──") or text == "---" or text == "***":
            continue

        is_heading2 = (
            "Heading 2" in style_name
            or style_name == "heading 2"
            or style_name == "Heading2"
        )

        if is_heading2:
            found_first_heading2 = True
            clean_title = _clean_chapter_title(text)

            if current_chapter and current_chapter.paragraphs:
                chapters.append(current_chapter)

            current_chapter = Chapter(title=clean_title)
            logger.debug("发现章节: %s", clean_title)
            continue

        if not found_first_heading2:
            summary_parts.append(text)
        elif current_chapter is not None:
            current_chapter.paragraphs.append(text)

    if current_chapter and current_chapter.paragraphs:
        chapters.append(current_chapter)

    summary = "\n".join(summary_parts).strip()

    logger.info(
        "本地 .docx 解析完成: %d 个章节, 摘要 %d 字, 正文 %d 字",
        len(chapters), len(summary), sum(ch.word_count for ch in chapters),
    )

    # 如果没有 Heading2，将整个文档作为单一章节
    if not chapters:
        logger.warning("未找到 Heading 2 章节分隔，将整个文档作为单一章节处理")
        all_text_parts = []
        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if text and not text.startswith("──"):
                all_text_parts.append(text)
        if all_text_parts:
            chapters = [Chapter(title=doc_title, paragraphs=all_text_parts)]

    # 清理临时文件
    Path(tmp_path).unlink(missing_ok=True)

    return DocumentContent(
        doc_title=doc_title,
        doc_token=file_token,
        summary=summary,
        chapters=chapters,
    )


def read_local_docx(docx_path: str | Path, doc_title: str | None = None) -> DocumentContent:
    """直接从本地 .docx 文件读取并按章节切分。

    适用于：飞书云盘下载失败时，直接读取本地已有的 .docx 文件。

    Args:
        docx_path: 本地 .docx 文件路径
        doc_title: 文档标题，如果不指定则从文件名推断

    Returns:
        解析后的 DocumentContent。
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("需要安装 python-docx 库: pip install python-docx")

    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"本地 .docx 文件不存在: {docx_path}")

    if not doc_title:
        doc_title = docx_path.stem  # 去掉 .docx 扩展名

    logger.info("从本地文件读取: %s（标题: %s）", docx_path, doc_title)
    docx_doc = DocxDocument(str(docx_path))

    # 解析段落
    summary_parts: list[str] = []
    chapters: list[Chapter] = []
    current_chapter: Chapter | None = None
    found_first_heading2 = False

    for para in docx_doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        if text.startswith("──") or text == "---" or text == "***":
            continue

        is_heading2 = (
            "Heading 2" in style_name
            or style_name == "heading 2"
            or style_name == "Heading2"
        )

        if is_heading2:
            found_first_heading2 = True
            clean_title = _clean_chapter_title(text)

            if current_chapter and current_chapter.paragraphs:
                chapters.append(current_chapter)

            current_chapter = Chapter(title=clean_title)
            logger.debug("发现章节: %s", clean_title)
            continue

        if not found_first_heading2:
            summary_parts.append(text)
        elif current_chapter is not None:
            current_chapter.paragraphs.append(text)

    if current_chapter and current_chapter.paragraphs:
        chapters.append(current_chapter)

    summary = "\n".join(summary_parts).strip()

    logger.info(
        "本地 .docx 解析完成: %d 个章节, 摘要 %d 字, 正文 %d 字",
        len(chapters), len(summary), sum(ch.word_count for ch in chapters),
    )

    if not chapters:
        logger.warning("未找到 Heading 2 章节分隔，将整个文档作为单一章节处理")
        all_text_parts = []
        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if text and not text.startswith("──"):
                all_text_parts.append(text)
        if all_text_parts:
            chapters = [Chapter(title=doc_title, paragraphs=all_text_parts)]

    return DocumentContent(
        doc_title=doc_title,
        doc_token=f"local:{docx_path.name}",
        summary=summary,
        chapters=chapters,
    )


def _clean_chapter_title(raw_title: str) -> str:
    """清理章节标题，去掉时间范围标记等装饰性内容。

    示例：
    - "第一章 投放基础（00:00 - 05:30）" → "第一章 投放基础"
    - "账户搭建 [00:00-05:30]" → "账户搭建"
    """
    # 移除常见的时间范围格式
    patterns = [
        r"[\(（]\s*\d{1,2}:\d{2}\s*[-–—]\s*\d{1,2}:\d{2}\s*[\)）]",  # (00:00 - 05:30)
        r"\[\s*\d{1,2}:\d{2}\s*[-–—]\s*\d{1,2}:\d{2}\s*\]",          # [00:00-05:30]
        r"\s*\d{1,2}:\d{2}\s*[-–—]\s*\d{1,2}:\d{2}\s*$",             # 末尾 00:00-05:30
    ]
    result = raw_title
    for pattern in patterns:
        result = re.sub(pattern, "", result)
    return result.strip()


# ── Obsidian Markdown 读取 ────────────────────────────────

def read_local_markdown(
    md_path: str | Path,
    doc_title: str | None = None,
) -> DocumentContent:
    """读取本地 Obsidian Markdown 文件并构建 DocumentContent。

    将整张 Playbook 卡片作为单一 Chapter 处理（不按粗体标记拆分），
    因为每张卡片本身已是完整主题单元（约 800-1200 字），体量适中。

    自动剥离 YAML frontmatter，提取 ``# 标题行`` 作为 doc_title。

    Args:
        md_path: 本地 .md 文件路径
        doc_title: 文档标题，如果不指定则从 ``# 标题行`` 或文件名推断

    Returns:
        解析后的 DocumentContent（单章节）。
    """
    md_path = Path(md_path)
    if not md_path.exists():
        raise FileNotFoundError(f"本地 .md 文件不存在: {md_path}")

    raw_text = md_path.read_text(encoding="utf-8")

    # ── 1. 剥离 YAML frontmatter ──────────────────────
    frontmatter: dict[str, Any] = {}
    body = raw_text
    fm_match = re.match(r"^---\s*\n(.*?)\n---\s*\n", raw_text, re.DOTALL)
    if fm_match:
        fm_text = fm_match.group(1)
        body = raw_text[fm_match.end():]
        # 简易 YAML 解析（避免引入 pyyaml 依赖）
        frontmatter = _parse_simple_yaml(fm_text)

    # ── 2. 提取 # 标题行 ──────────────────────────────
    h1_match = re.search(r"^#\s+(.+)$", body, re.MULTILINE)
    if not doc_title:
        if h1_match:
            doc_title = h1_match.group(1).strip()
        else:
            doc_title = md_path.stem  # 去掉 .md 扩展名

    # ── 3. 构建 summary（从 frontmatter 拼接） ────────
    summary_parts = []
    if frontmatter.get("tags"):
        tags = frontmatter["tags"]
        if isinstance(tags, list):
            summary_parts.append("标签: " + ", ".join(str(t) for t in tags))
        else:
            summary_parts.append(f"标签: {tags}")
    if frontmatter.get("industry_focus"):
        focus = frontmatter["industry_focus"]
        if isinstance(focus, list):
            summary_parts.append("行业: " + ", ".join(str(f) for f in focus))
        else:
            summary_parts.append(f"行业: {focus}")
    if frontmatter.get("inspiration"):
        insp = frontmatter["inspiration"]
        if isinstance(insp, list):
            summary_parts.append("灵感案例: " + ", ".join(str(i) for i in insp))
        else:
            summary_parts.append(f"灵感案例: {insp}")
    summary = " | ".join(summary_parts) if summary_parts else ""

    # ── 4. 整张卡片作为单一 Chapter ───────────────────
    # 正文 = 去掉 # 标题行后的全部内容
    content_body = body
    if h1_match:
        # 去掉 # 标题行那一行
        content_body = body[:h1_match.start()] + body[h1_match.end():]
    content_body = content_body.strip()

    chapter = Chapter(
        title=md_path.stem,  # 文件名作为章节标题
        paragraphs=[content_body] if content_body else [],
    )

    logger.info(
        "Obsidian MD 解析完成: %s — 1 章节, %d 字",
        doc_title, chapter.word_count,
    )

    return DocumentContent(
        doc_title=doc_title,
        doc_token=f"local-md:{md_path.name}",
        summary=summary,
        chapters=[chapter] if chapter.paragraphs else [],
    )


def _parse_simple_yaml(text: str) -> dict[str, Any]:
    """极简 YAML 解析器，只处理 key: value 和 key: [list] 格式。

    不引入 pyyaml 依赖。支持：
    - ``key: value``
    - ``key: [a, b, c]``
    """
    result: dict[str, Any] = {}
    for line in text.split("\n"):
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        match = re.match(r"^(\w[\w_-]*)\s*:\s*(.*)$", line)
        if not match:
            continue
        key = match.group(1)
        val = match.group(2).strip()
        # 尝试解析列表 [a, b, c]
        list_match = re.match(r"^\[(.*)\]$", val)
        if list_match:
            items = [item.strip().strip("'\"") for item in list_match.group(1).split(",")]
            result[key] = [i for i in items if i]
        else:
            result[key] = val.strip("'\"")
    return result
